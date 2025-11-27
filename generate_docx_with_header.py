#!/usr/bin/env python3.11
"""
Script pour générer des documents Word avec le papier à en-tête PORTEO GROUP
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import markdown
import re
import sys

def copy_header_from_template(doc, template_path):
    """Copie l'en-tête du template PORTEO vers le nouveau document"""
    template = Document(template_path)
    
    # Copier l'en-tête de la première section du template
    template_header = template.sections[0].header
    new_header = doc.sections[0].header
    
    # Copier tous les paragraphes de l'en-tête
    for para in template_header.paragraphs:
        new_para = new_header.add_paragraph()
        new_para.alignment = para.alignment
        
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
            
            # Copier les images
            for drawing in run._element.findall(qn('w:drawing')):
                # Copier l'élément drawing complet
                new_run._element.append(drawing)
    
    # Copier les relations d'images
    for rel_id, rel in template_header.part.rels.items():
        if "image" in rel.reltype:
            # Ajouter la relation d'image au nouvel en-tête
            new_header.part.relate_to(rel.target_part, rel.reltype)

def markdown_to_docx(markdown_text, output_path, template_path):
    """
    Convertit du texte Markdown en document Word avec papier à en-tête PORTEO
    
    Args:
        markdown_text: Texte au format Markdown
        output_path: Chemin du fichier .docx à créer
        template_path: Chemin du template avec papier à en-tête
    """
    # Créer un nouveau document
    doc = Document()
    
    # Copier l'en-tête du template
    copy_header_from_template(doc, template_path)
    
    # Configuration des marges (en inches)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)  # Plus d'espace pour l'en-tête
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parser le Markdown ligne par ligne
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Ligne vide
            doc.add_paragraph()
            i += 1
            continue
        
        # Titres
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        
        # Listes à puces
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Listes numérotées
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Citations
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:], style='Intense Quote')
        
        # Paragraphe normal
        else:
            p = doc.add_paragraph()
            
            # Traiter le texte avec formatage inline (gras, italique)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Gras
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    # Italique
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    # Code inline
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)
        
        i += 1
    
    # Sauvegarder le document
    doc.save(output_path)
    print(f"Document créé: {output_path}")

def main():
    if len(sys.argv) < 4:
        print("Usage: python generate_docx_with_header.py <markdown_file> <output_file> <template_file>")
        sys.exit(1)
    
    markdown_file = sys.argv[1]
    output_file = sys.argv[2]
    template_file = sys.argv[3]
    
    # Lire le contenu Markdown
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # Générer le document Word
    markdown_to_docx(markdown_text, output_file, template_file)

if __name__ == "__main__":
    main()
