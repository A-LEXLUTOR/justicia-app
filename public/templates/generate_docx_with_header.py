#!/usr/bin/env python3.11
"""
Script pour générer des documents Word avec le papier à en-tête PORTEO GROUP
Avec couleur #BA8A52 et typographie Just Sans Variable
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import markdown
import re
import sys

def copy_header_from_template(doc, template_path):
    """Copie l'en-tête du template PORTEO vers le nouveau document"""
    template = Document(template_path)
    
    # Copier l'en-tête de la première section du template
    template_header = template.sections[0].header
    new_header = doc.sections[0].header
    
    # Copier tous les paragraphes de l'en-tête
    for para in template_header.paragraphs:
        new_para = new_header.add_paragraph()
        new_para.alignment = para.alignment
        
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
            
            # Copier les images
            for drawing in run._element.findall(qn('w:drawing')):
                # Copier l'élément drawing complet
                new_run._element.append(drawing)
    
    # Copier les relations d'images
    for rel_id, rel in template_header.part.rels.items():
        if "image" in rel.reltype:
            # Ajouter la relation d'image au nouvel en-tête
            new_header.part.relate_to(rel.target_part, rel.reltype)

def remove_ai_introduction(text):
    """
    Supprime les phrases d'introduction de l'IA
    """
    # Patterns à supprimer
    patterns = [
        r"^Voici.*?:\s*",
        r"^Voici.*?\.\s*",
        r"^Je.*?:\s*",
        r"^Ci-dessous.*?:\s*",
        r"^Voilà.*?:\s*",
    ]
    
    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.MULTILINE | re.DOTALL)
    
    return text.strip()

def markdown_to_docx(markdown_text, output_path, template_path):
    """
    Convertit du texte Markdown en document Word avec papier à en-tête PORTEO
    """
    # Supprimer le texte d'introduction de l'IA
    markdown_text = remove_ai_introduction(markdown_text)
    
    # Créer un nouveau document
    doc = Document()
    
    # Copier l'en-tête du template
    copy_header_from_template(doc, template_path)
    
    # Couleur PORTEO (#BA8A52)
    porteo_color = RGBColor(186, 138, 82)
    
    # Traiter le texte ligne par ligne
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Ligne vide
            doc.add_paragraph()
            continue
        
        # Titre niveau 1 (# Titre)
        if line.startswith('# '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = 'Just Sans Variable ExtraBold'
            run.font.size = Pt(50)
            run.font.color.rgb = porteo_color
            run.bold = True
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        # Titre niveau 2 (## Titre)
        elif line.startswith('## '):
            text = line[3:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = 'Just Sans Variable Regular'
            run.font.size = Pt(30)
            run.font.color.rgb = porteo_color
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        # Titre niveau 3 (### Titre)
        elif line.startswith('### '):
            text = line[4:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = 'Just Sans Variable Bold'
            run.font.size = Pt(20)
            run.font.color.rgb = porteo_color
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        # Liste à puces
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            para = doc.add_paragraph(text, style='List Bullet')
            for run in para.runs:
                run.font.name = 'Just Sans Variable Regular'
                run.font.size = Pt(16)
            
        # Liste numérotée
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            para = doc.add_paragraph(text, style='List Number')
            for run in para.runs:
                run.font.name = 'Just Sans Variable Regular'
                run.font.size = Pt(16)
            
        # Citation (> texte)
        elif line.startswith('> '):
            text = line[2:].strip()
            para = doc.add_paragraph(text)
            para.style = 'Quote'
            for run in para.runs:
                run.font.name = 'Just Sans Variable Regular'
                run.font.size = Pt(16)
                run.italic = True
            
        # Paragraphe normal
        else:
            # Traiter le formatage inline (gras, italique)
            para = doc.add_paragraph()
            
            # Remplacer **texte** par gras
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    text = part[2:-2]
                    run = para.add_run(text)
                    run.bold = True
                    run.font.name = 'Just Sans Variable Bold'
                    run.font.size = Pt(16)
                elif part.startswith('*') and part.endswith('*'):
                    text = part[1:-1]
                    run = para.add_run(text)
                    run.italic = True
                    run.font.name = 'Just Sans Variable Regular'
                    run.font.size = Pt(16)
                else:
                    run = para.add_run(part)
                    run.font.name = 'Just Sans Variable Regular'
                    run.font.size = Pt(16)
    
    # Sauvegarder le document
    doc.save(output_path)
    print(f"Document créé : {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python3.11 generate_docx_with_header.py <input.md> <output.docx> <template.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    template_file = sys.argv[3]
    
    # Lire le fichier Markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # Générer le document Word
    markdown_to_docx(markdown_text, output_file, template_file)
